# -*- coding: utf-8 -*-
"""The documents students hand in.

One definition per document, in one place, so every hand-in a student
touches has the same header, the same identity block, the same heading
weights and the same fill-in boxes. Before this, ten documents had been
written at different times by different hands and it showed: some opened
with a bare "Name" on line one, some had the title in the body, some
asked for the project title and some did not, and two of them were the
same instrument under two different names.

WHAT WAS DELIBERATELY CHANGED, and why -- so nobody "fixes" it back:

  * Logbook and Daily Journal were the same document under two names, both
    being distributed to the same class at the same time. They are now one
    document, "Daily Logbook". Whichever one Classroom is handing out, it
    should now hand out this.
  * Four templates were 3.36 MB each because of a full-resolution image.
    That is roughly a thousand times what a text template should weigh, and
    every student copy carried it, every year, in every class. The logo here
    is 64 KB and looks the same on paper.
  * The instruction lines are addressed to the student in plain words.
    "Utilize the following status codes to categorize the current phase of
    engineering activity" says the same thing as "mark which stage of the
    design process you were in", and one of them is readable at 7:45 am.
  * Every document now has the same identity block, so a hand-in is never
    anonymous and never undated.

WHAT WAS NOT CHANGED: the questions themselves. Every prompt below is the
prompt from the existing template, trimmed rather than reworded, because
the wording is Dan's and the questions are the assessment.

FIELD TYPES
  ('h',    text)                     a section heading
  ('note', text)                     one italic line of guidance under it
  ('box',  n_lines)                  a bordered box the student types into
  ('bul',  [placeholders])           a bulleted list to extend
  ('tbl',  [headers], n_rows)        a table to fill in
  ('rule',)                          a hairline separator
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(HERE, 'logo', 'header.png')

# --- the house look: the SITE's design system, on paper --------------------
# Same three faces and same palette as bhr-shop-hub, so a student moving
# between the website and a hand-in sees one thing, not two. All three faces
# ship with Google Docs, which is where these get filled in, so naming them
# is enough -- no font file travels with the document.
#
#   display   Chakra Petch      site headings
#   prose     Source Serif 4    site body text and the italic guidance
#   mono      Space Mono        site eyebrows, section heads, codes
#
# No page background. The site's paper is a near-white; on a document that
# prints and gets marked up, plain white is the honest equivalent.
FACE_DISPLAY = 'Chakra Petch'
FACE_BODY    = 'Chakra Petch'
FACE_PROSE   = 'Source Serif 4'
FACE_MONO    = 'Space Mono'
FACE_HEAD    = 'Space Mono'

PURPLE = RGBColor(0x6b, 0x47, 0x85)   # --accent
INK    = RGBColor(0x26, 0x2b, 0x39)   # --ink
INK2   = RGBColor(0x4f, 0x55, 0x66)   # --ink-2
INK3   = RGBColor(0x7d, 0x83, 0x94)   # --ink-3
RULE   = 'dedbd5'                     # --rule
BOXRULE = 'b6b1bf'                    # a touch stronger, so a fill-in box reads
SOFT   = 'f1ebf7'                     # --accent-soft

# the design-process palette, kept semantic but pulled toward the site
EDP = [('PI',  'Problem identification / research',  RGBColor(0x1d, 0x4e, 0x89)),
       ('DD',  'Detailed design (CAD / modelling)',  RGBColor(0x8a, 0x64, 0x10)),
       ('FAB', 'Fabrication & development',          RGBColor(0x2f, 0x6b, 0x3a)),
       ('TE',  'Testing & evaluation',               RGBColor(0x99, 0x33, 0x2b)),
       ('IR',  'Improve & redesign',                 RGBColor(0x6b, 0x47, 0x85))]

# How tall a fill-in box is IS the instruction. A student reads box height
# before they read the prompt, and a three-word answer in a half-page box
# feels wrong to them, so they pad it. Sizes are named rather than numeric so
# the same question type gets the same box in every document.
BOX = {
    'xs': 1,    # a word, a number, a date
    's':  2,    # a sentence or two
    'm':  4,    # a paragraph
    'l':  7,    # a real piece of writing
    'xl': 11,   # a full reflection, or an area to paste images into
}

TEXT_W = 6.8                  # 8.5in page less two 0.85in margins
TEXT_TWIPS = int(TEXT_W * 1440)
TEXT_EMU = int(TEXT_W * 914400)

SCHOOL = 'Blue Hills Regional Technical School  ·  Room E-126'

# what to show in the value cell of the identity block, so the student is
# never guessing what a field wants
IDENT_HINT = {
    'Name': '[Your name]',
    'Project engineer': '[Your name]',
    'Date': '[Insert date]',
    'Project': '[Project title]',
    'Project(s)': '[List all active projects]',
    'Class': '[Grade and course]',
    'Term': '[Term]',
    'Tools and software used': '[Hardware and software you used]',
    'Instructor': '[Mr. Frank or Mr. Dryer]',
    'Pathway': '[Which pathway]',
}


# ------------------------------------------------------------------ helpers

def _shade(cell, hexfill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def _borders(table, colour=RULE, size=4, inside=True):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    edges = ['top', 'left', 'bottom', 'right']
    if inside:
        edges += ['insideH', 'insideV']
    for edge in edges:
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(size))
        e.set(qn('w:color'), colour)
        borders.append(e)
    tblPr.append(borders)


def _full_width(table):
    """python-docx leaves a one-column table at its default width, which on
    the page reads as a box that stopped short. Force every fill-in area to
    the full text column."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    w = OxmlElement('w:tblW')
    w.set(qn('w:type'), 'dxa')
    w.set(qn('w:w'), str(TEXT_TWIPS))
    tblPr.append(w)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    # Word and LibreOffice both prefer explicit cell widths over the table
    # width when the layout is fixed, so set them too.
    n = len(table.columns)
    for col in table.columns:
        for c in col.cells:
            c.width = Emu(int(TEXT_EMU / n))


def _no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def _run(p, text, size=10.5, bold=False, italic=False, colour=INK,
         font=FACE_BODY):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = colour
    r.font.name = font
    return r


def _para(container, text='', size=10.5, bold=False, italic=False,
          colour=INK, before=0, after=4, align=None):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        _run(p, text, size, bold, italic, colour)
    return p


# ------------------------------------------------------------------- chrome

def _page_background(doc, hexfill):
    """A full-page colour wash. Word needs displayBackgroundShape set in
    settings.xml or it ignores w:background entirely; Google Docs honours
    both."""
    el = OxmlElement('w:background')
    el.set(qn('w:color'), hexfill)
    doc.element.insert(0, el)
    st = doc.settings.element
    d = OxmlElement('w:displayBackgroundShape')
    st.insert(0, d)


def _page_setup(doc):
    s = doc.sections[0]
    s.top_margin = Inches(0.85)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.85)
    s.right_margin = Inches(0.85)
    s.header_distance = Inches(0.35)
    s.footer_distance = Inches(0.3)

    for name, sz in (('Heading 2', 13), ('Heading 3', 11), ('Heading 4', 10.5)):
        st = doc.styles[name]
        st.font.name = FACE_HEAD if name != 'Heading 4' else FACE_BODY
        st.font.size = Pt(sz)
        st.font.color.rgb = INK
        st.font.bold = True
        st.font.italic = False

    normal = doc.styles['Normal']
    normal.font.name = FACE_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)


def _header(doc):
    hdr = doc.sections[0].header
    t = hdr.add_table(1, 2, Inches(6.8))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(t)
    t.columns[0].width = Inches(4.4)
    t.columns[1].width = Inches(2.4)

    left = t.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO):
        left.add_run().add_picture(LOGO, height=Inches(0.30))
        _run(left, '  ', 9)
    _run(left, 'BHR ENGINEERING TECHNOLOGY', 8.5, bold=True, colour=PURPLE)

    right = t.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    _run(right, SCHOOL, 7.5, colour=INK3)

    # hairline under the header
    p = hdr.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '8')
    b.set(qn('w:color'), '141c26')
    pbdr.append(b)
    pPr.append(pbdr)


def _footer(doc, docname):
    ftr = doc.sections[0].footer
    p = ftr.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    _run(p, docname, 7.5, colour=INK3)
    _run(p, '\t\t', 7.5)
    _run(p, 'Page ', 7.5, colour=INK3)
    # PAGE field
    r = p.add_run()
    r.font.size = Pt(7.5)
    r.font.color.rgb = INK3
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    r._r.addnext(fld)


def _title(doc, title, standfirst):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    _run(p, title, size=21, bold=True, colour=INK, font=FACE_DISPLAY)
    if standfirst:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(11)
        _run(q, standfirst, size=10, italic=True, colour=INK2,
             font=FACE_PROSE)


def _identity(doc, fields):
    """The block at the top of every hand-in: who, when, what."""
    rows = (len(fields) + 1) // 2
    t = doc.add_table(rows=rows, cols=4)
    _borders(t, colour=BOXRULE, size=6)
    _full_width(t)
    for i, label in enumerate(fields):
        c_label = t.cell(i // 2, (i % 2) * 2)
        c_value = t.cell(i // 2, (i % 2) * 2 + 1)
        _shade(c_label, SOFT)
        p = c_label.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        _run(p, label, 9, bold=True, colour=PURPLE)
        vp = c_value.paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        vp.paragraph_format.space_before = Pt(2)
        _run(vp, IDENT_HINT.get(label, '[\u2026]'), 9.5, colour=INK3)
    for i, w in enumerate((1.15, 2.25, 1.15, 2.25)):
        for cell in t.columns[i].cells:
            cell.width = Inches(w)
    _para(doc, '', after=6)


PURPLE_HEX = '6b4785'


def _pbdr(p, colour, size):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(size))
    b.set(qn('w:space'), '3')
    b.set(qn('w:color'), colour)
    pbdr.append(b)
    pPr.append(pbdr)


def _dropdown(p, options, default=None, width_chars=6):
    """A real dropdown the student picks from, as a Word content control.

    Word (desktop and web) renders this as a clickable list. Google Docs does
    not have an importer for content controls, so on upload it flattens to the
    default value as plain text -- which is why the default is a real, sensible
    option rather than a blank. If these templates are going to live as Google
    Docs, add the native Docs dropdown once on the master copy; every student
    copy made from it inherits the dropdown.
    """
    sdt = OxmlElement('w:sdt')
    pr = OxmlElement('w:sdtPr')

    alias = OxmlElement('w:alias')
    alias.set(qn('w:val'), 'Status')
    pr.append(alias)

    rpr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), FACE_MONO)
    rf.set(qn('w:hAnsi'), FACE_MONO)
    rpr.append(rf)
    b = OxmlElement('w:b')
    rpr.append(b)
    pr.append(rpr)

    ddl = OxmlElement('w:dropDownList')
    for o in options:
        li = OxmlElement('w:listItem')
        li.set(qn('w:displayText'), o)
        li.set(qn('w:value'), o)
        ddl.append(li)
    pr.append(ddl)
    sdt.append(pr)

    content = OxmlElement('w:sdtContent')
    r = OxmlElement('w:r')
    rp = OxmlElement('w:rPr')
    rf2 = OxmlElement('w:rFonts')
    rf2.set(qn('w:ascii'), FACE_MONO)
    rf2.set(qn('w:hAnsi'), FACE_MONO)
    rp.append(rf2)
    bb = OxmlElement('w:b')
    rp.append(bb)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '19')
    rp.append(sz)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FFFFFF')
    rp.append(shd)
    r.append(rp)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = ' %s ' % (default or options[0])
    r.append(t)
    content.append(r)
    sdt.append(content)
    p._p.append(sdt)


# ------------------------------------------------------------------- fields

def _emit(doc, item):
    kind = item[0]

    if kind == 'h':
        # Source Code Pro Black. This is the beat that makes the page read as
        # an engineering document rather than a worksheet: a heavy monospace
        # header with the interval name at a smaller size beside it.
        # A REAL heading style, not a formatted paragraph. This is what makes
        # the section collapsible in Google Docs -- the caret only appears on
        # built-in Heading 1-6. Run-level formatting below overrides the
        # style's own look, so it collapses AND matches the site.
        p = doc.add_paragraph(style='Heading 2')
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        txt = item[1]
        if ':' in txt:
            head, rest = txt.split(':', 1)
            _run(p, head.upper() + ': ', size=13, bold=True, colour=PURPLE,
                 font=FACE_HEAD)
            _run(p, rest.strip(), size=11, bold=True, colour=INK,
                 font=FACE_HEAD)
        else:
            _run(p, txt.upper(), size=13, bold=True, colour=PURPLE,
                 font=FACE_HEAD)
        # the site draws a 2px accent rule under every term heading
        _pbdr(p, PURPLE_HEX, 12)
        # the status-code chip sits on the same line, small, so the heading
        # never wraps because of it
        if len(item) > 2 and item[2]:
            _run(p, '     EDP STATUS CODE: ', size=9, bold=True, colour=INK)
            _dropdown(p, [c for c, _n, _col in EDP], default=item[2])

    elif kind == 'h2':
        p = doc.add_paragraph(style='Heading 3')
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(2)
        _run(p, item[1].upper(), size=11, bold=True, colour=PURPLE,
             font=FACE_HEAD)

    elif kind == 'pick':
        # label + dropdown on one line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        _run(p, item[1] + ': ', size=10.5, bold=True, colour=INK)
        _dropdown(p, item[2], default=item[3] if len(item) > 3 else None)
        if len(item) > 4 and item[4]:
            _run(p, '   ', size=10.5)
            _run(p, item[4], size=9.5, italic=True, colour=INK2,
                 font=FACE_PROSE)

    elif kind == 'label':
        # Bold label with its guidance italic on the SAME line -- the original
        # does this everywhere and it is most of why the page reads dense and
        # professional rather than like a form. Heading 4 so it collapses and
        # nests under its interval in the Docs outline.
        p = doc.add_paragraph(style='Heading 4')
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        _run(p, item[1], size=10.5, bold=True, colour=INK)
        if len(item) > 2 and item[2]:
            _run(p, '  \u2014  ', size=10.5, colour=INK3)
            _run(p, item[2], size=9.5, italic=True, colour=INK2,
                 font=FACE_PROSE)

    elif kind == 'edp':
        # the colour-coded design-process legend
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        for i, (code, name, col) in enumerate(EDP):
            if i:
                _run(p, '   ', size=10, colour=INK3, font=FACE_MONO)
            _run(p, '[%s] ' % code, size=9, bold=True, colour=col,
                 font=FACE_MONO)
            _run(p, name, size=9, colour=col, font=FACE_PROSE)

    elif kind == 'note':
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(5)
        _run(q, item[1], size=9.5, italic=True, colour=INK2, font=FACE_PROSE)

    elif kind == 'box':
        t = doc.add_table(rows=1, cols=1)
        _borders(t, colour=BOXRULE, size=6)
        _full_width(t)
        cell = t.cell(0, 0)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        # a placeholder, so a student on a laptop can see where to click and
        # what is wanted. An empty box tells them nothing.
        _run(p, item[2] if len(item) > 2 else '[Type here\u2026]',
             10.5, colour=INK3)
        nlines = BOX.get(item[1], item[1]) if isinstance(item[1], str) \
            else item[1]
        for _ in range(max(0, nlines - 1)):
            q = cell.add_paragraph()
            q.paragraph_format.space_after = Pt(3)
        _para(doc, '', after=2)

    elif kind == 'bul':
        # bulleted placeholders inside a bordered box, as the original does
        t = doc.add_table(rows=1, cols=1)
        _borders(t, colour=BOXRULE, size=6)
        _full_width(t)
        cell = t.cell(0, 0)
        first = True
        for text in item[1]:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.style = doc.styles['List Bullet']
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _run(p, text, 10.5, colour=INK3)
        _para(doc, '', after=2)

    elif kind == 'tbl':
        heads, nrows = item[1], item[2]
        t = doc.add_table(rows=nrows + 1, cols=len(heads))
        _borders(t, colour=BOXRULE, size=6)
        _full_width(t)
        for i, h in enumerate(heads):
            c = t.cell(0, i)
            _shade(c, SOFT)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _run(p, h, 8.5, bold=True, colour=PURPLE)
        for r in range(1, nrows + 1):
            for i in range(len(heads)):
                pp = t.cell(r, i).paragraphs[0]
                pp.paragraph_format.space_before = Pt(3)
                pp.paragraph_format.space_after = Pt(3)
        _para(doc, '', after=2)

    elif kind == 'rule':
        _para(doc, '', after=2)


def build(spec, outdir):
    doc = Document()
    _page_setup(doc)
    _header(doc)
    _footer(doc, spec['name'])
    _title(doc, spec['title'], spec.get('standfirst', ''))
    _identity(doc, spec['identity'])
    for item in spec['body']:
        _emit(doc, item)
    path = os.path.join(outdir, spec['file'])
    doc.save(path)
    return path
